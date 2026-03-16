from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def set_chinese_font(run, font_name='SimSun', size=11, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, 'SimHei', 16, True)
        run.font.color.rgb = RGBColor(0, 0, 128)
    elif level == 2:
        set_chinese_font(run, 'SimHei', 14, True)
    else:
        set_chinese_font(run, 'SimHei', 12, True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', size, bold)
    return p

def add_table_zh(doc, rows, cols, style='Table Grid'):
    """添加表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    return table

# 创建文档
doc = Document()

# 设置文档默认字体
doc.styles['Normal'].font.name = 'SimSun'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ========== 封面 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Horizon Watches\n运营手册')
set_chinese_font(run, 'SimHei', 28, True)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
now = datetime.now()
date_time_str = now.strftime('版本 2.2 | %Y年%m月%d日 %H:%M 生成')
run = subtitle.add_run(date_time_str)
set_chinese_font(run, 'SimSun', 14)

doc.add_page_break()

# ========== 目录 ==========
add_heading_zh(doc, '目录', 1)
toc_items = [
    '1. 项目概述',
    '2. 系统架构',
    '3. 后台管理操作指南',
    '4. 钱包管理与资金流转',
    '5. 订单处理流程',
    '6. 数据库配置说明',
    '7. 常见问题与故障处理',
    '8. 附录：账号信息'
]
for item in toc_items:
    add_paragraph_zh(doc, item)

doc.add_page_break()

# ========== 第1章：项目概述 ==========
add_heading_zh(doc, '1. 项目概述', 1)

add_heading_zh(doc, '1.1 项目简介', 2)
add_paragraph_zh(doc, 'Horizon Watches 是一个面向欧洲市场的独立站电商项目，主营复古风格腕表。采用 D2C（Direct to Consumer）模式，通过独立站直接销售给客户。')

add_heading_zh(doc, '1.2 技术栈', 2)
tech_stack = [
    '前端：Next.js 14 + TypeScript + Tailwind CSS',
    '部署：Vercel + Surge.sh（双部署）',
    '数据库：PostgreSQL（Supabase）',
    '支付：USDT (TRC20)',
    '客服：AI ChatBot + 人工兜底'
]
for item in tech_stack:
    add_paragraph_zh(doc, '• ' + item)

add_heading_zh(doc, '1.3 核心网址', 2)
urls = [
    '主站：https://horizon-watch.vercel.app',
    '管理后台：https://horizon-watch-admin.vercel.app/admin'
]
for url in urls:
    add_paragraph_zh(doc, '• ' + url)

doc.add_page_break()

# ========== 第2章：系统架构 ==========
add_heading_zh(doc, '2. 系统架构', 1)

add_heading_zh(doc, '2.1 多角色权限体系', 2)

table = add_table_zh(doc, 5, 3)
headers = ['角色', '权限范围', '使用场景']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

roles_data = [
    ['SUPERADMIN\n超级管理员', '商品管理、订单管理、用户管理、钱包配置、财务设置', 'S1使用，全权管理'],
    ['ADMIN\n管理员', '商品管理、订单管理', '综合运营人员'],
    ['SUPPLY\n供应链', '商品管理、库存管理', '负责商品和库存'],
    ['LOGISTICS\n物流', '订单管理、发货处理', '负责订单发货']
]

for i, row_data in enumerate(roles_data, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 10)

add_heading_zh(doc, '2.2 默认账号信息', 2)

table = add_table_zh(doc, 5, 4)
headers = ['用户名', '密码', '角色', '邮箱']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

accounts = [
    ['superadmin', 'super123', 'SUPERADMIN', 'superadmin@horizonwatches.com'],
    ['admin', 'admin123', 'ADMIN', 'admin@horizonwatches.com'],
    ['supply', 'supply123', 'SUPPLY', 'supply@horizonwatches.com'],
    ['logistics', 'logistics123', 'LOGISTICS', 'logistics@horizonwatches.com']
]

for i, row_data in enumerate(accounts, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 10)
        if j == 1:  # 密码列加粗
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# ========== 第3章：后台管理操作指南 ==========
add_heading_zh(doc, '3. 后台管理操作指南', 1)

add_heading_zh(doc, '3.1 登录后台', 2)
steps = [
    '访问后台地址：https://horizon-watch-store-1773050228.surge.sh/admin',
    '输入用户名和密码（见第2章账号信息）',
    '点击 Login 登录',
    '登录后根据角色显示对应功能菜单'
]
for i, step in enumerate(steps, 1):
    add_paragraph_zh(doc, f'{i}. {step}')

add_heading_zh(doc, '3.2 商品管理（Products）', 2)
add_paragraph_zh(doc, '功能说明：添加、编辑、删除商品，管理库存。', True)

features = [
    '查看商品列表：显示所有商品，支持搜索',
    '添加商品：点击 "Add Product" 按钮，填写商品信息',
    '编辑商品：点击商品行末的编辑图标',
    '删除商品：点击商品行末的删除图标（需确认）',
    '导入/导出：支持JSON格式批量导入导出'
]
for f in features:
    add_paragraph_zh(doc, '• ' + f)

add_heading_zh(doc, '3.3 订单管理（Orders）', 2)
add_paragraph_zh(doc, '功能说明：查看订单、更新订单状态、录入物流信息。', True)
add_paragraph_zh(doc, '注：完整订单功能需要连接数据库后生效。')

add_heading_zh(doc, '3.4 用户管理（Users）- SUPERADMIN专属', 2)
add_paragraph_zh(doc, '功能说明：管理系统账号，添加/编辑/删除用户。', True)

user_mgmt = [
    '查看用户列表：显示所有账号的用户名、角色、权限',
    '添加用户：点击 "Add User"，输入用户名、密码、选择角色',
    '编辑用户：点击编辑图标，可修改角色（用户名不可改）',
    '删除用户：点击删除图标，确认后删除',
    '角色参考：页面底部显示各角色权限对照表'
]
for f in user_mgmt:
    add_paragraph_zh(doc, '• ' + f)

add_heading_zh(doc, '3.5 财务设置（Settings）- SUPERADMIN专属', 2)
add_paragraph_zh(doc, '功能说明：配置3级USDT钱包地址。', True)

doc.add_page_break()

# ========== 第4章：钱包管理与资金流转 ==========
add_heading_zh(doc, '4. 钱包管理与资金流转', 1)

add_heading_zh(doc, '4.1 3级钱包架构', 2)

wallet_arch = '''
客户付款
    ↓
Level 1: 收款钱包 (仅接收)
    TGPBhfjSuwjrfGUtdqt6EZUbzhbRCGfC5c
    ↓ 每日归集
Level 2: 运营钱包 (支出用途)
    TCWgr2qGcheRsD7kceoFpJfMg59fFrJGCq
    ↓ 每周/每月汇总
Level 3: 利润钱包 (最终归属)
    TUyTqV47pd7o3Bg6Uhw5XJ9rwkdgi6tsKb ← STAR持有
'''
add_paragraph_zh(doc, wallet_arch)

add_heading_zh(doc, '4.2 钱包配置表', 2)

table = add_table_zh(doc, 4, 4)
headers = ['层级', '钱包地址', '用途', '操作权限']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

wallet_data = [
    ['L1 收款', 'TGPBhfjSuwjrfGUtdqt6EZUbzhbRCGfC5c', '仅用于网站展示，接收客户付款', 'S1查看，不主动操作'],
    ['L2 运营', 'TCWgr2qGcheRsD7kceoFpJfMg59fFrJGCq', '广告费、物流费、网站运维、应急资金', 'S1全权管理'],
    ['L3 利润', 'TUyTqV47pd7o3Bg6Uhw5XJ9rwkdgi6tsKb', '最终利润归集，定期提现', 'STAR持有']
]

for i, row_data in enumerate(wallet_data, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_heading_zh(doc, '4.3 资金流转规则', 2)
add_paragraph_zh(doc, '收入分成比例：', True)
add_paragraph_zh(doc, '• 40% 再投入：从L2转入广告/运营支出')
add_paragraph_zh(doc, '• 60% 利润：从L1直接转入L3（或经L2汇总）')

add_heading_zh(doc, '4.4 每日操作流程', 2)

add_paragraph_zh(doc, '早上 (9:00 AM)', True)
morning_tasks = [
    '检查L1钱包新到账',
    '核对订单与付款匹配',
    '更新订单状态',
    '准备发货清单给STAR'
]
for t in morning_tasks:
    add_paragraph_zh(doc, '  [ ] ' + t)

add_paragraph_zh(doc, '下午 (6:00 PM)', True)
evening_tasks = [
    '汇总当日订单',
    '执行资金归集（L1→L2/L3）',
    '记录财务流水'
]
for t in evening_tasks:
    add_paragraph_zh(doc, '  [ ] ' + t)

add_heading_zh(doc, '4.5 财务记录模板', 2)

table = add_table_zh(doc, 4, 7)
headers = ['日期', '订单号', '金额(USDT)', 'L1余额', '转入L2', '转入L3', '备注']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 10, True)

sample_data = [
    ['3/15', '#1001', '1,350', '5,680', '540', '810', 'Daytona'],
    ['3/15', '#1002', '1,165', '6,845', '466', '699', 'Submariner'],
    ['3/15', '汇总', '2,515', '-', '1,006', '1,509', '40%/60%']
]
for i, row_data in enumerate(sample_data, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

doc.add_page_break()

# ========== 第5章：订单处理流程 ==========
add_heading_zh(doc, '5. 订单处理流程', 1)

add_heading_zh(doc, '5.1 订单状态流转', 2)

status_flow = '''
待付款 → 已付款 → 待发货 → 已发货 → 已完成
   ↓
已取消 (超时未付款/客户取消)
'''
add_paragraph_zh(doc, status_flow)

add_heading_zh(doc, '5.2 完整订单处理流程', 2)

flow_steps = [
    ('客户下单', '客户在网站选择商品，填写收货信息，选择USDT支付'),
    ('客户付款', '客户转账到L1钱包，点击"已完成付款"'),
    ('收款确认', 'S1检查L1到账，核对金额与订单匹配，确认收款'),
    ('订单确认', '系统更新订单状态为"已付款"'),
    ('准备发货', 'S1生成发货清单，交给STAR安排物流'),
    ('发货处理', 'STAR打包发货，提供物流单号'),
    ('录入物流', 'S1在后台录入物流公司和单号，状态变为"已发货"'),
    ('资金归集', '按40%/60%比例将资金从L1转入L2和L3'),
    ('订单完成', '客户确认收货，订单完成')
]

for i, (title, desc) in enumerate(flow_steps, 1):
    add_paragraph_zh(doc, f'{i}. {title}', True)
    add_paragraph_zh(doc, f'   {desc}')

add_heading_zh(doc, '5.3 物流信息', 2)
logistics_info = [
    '物流渠道：欧洲敏感专线（17国覆盖）',
    '时效：9-15工作日',
    '单只手表运费：€5-8',
    '承运商：敏感专线-欧美-通'
]
for info in logistics_info:
    add_paragraph_zh(doc, '• ' + info)

doc.add_page_break()

# ========== 第6章：数据库配置说明 ==========
add_heading_zh(doc, '6. 数据库配置说明', 1)

add_heading_zh(doc, '6.1 当前状态', 2)
add_paragraph_zh(doc, '✅ 数据库架构设计：已完成')
add_paragraph_zh(doc, '⏳ 实际部署：等待 STAR 选择服务商后执行')

add_heading_zh(doc, '6.2 已完成内容', 2)
completed_items = [
    '数据库Schema设计（PostgreSQL）',
    'Prisma ORM配置',
    '后台API接口（/api/admin/*）',
    '数据导出功能'
]
for item in completed_items:
    add_paragraph_zh(doc, '• ' + item)

add_heading_zh(doc, '6.3 待决策事项', 2)
add_paragraph_zh(doc, '需要 STAR 选择数据库服务商：', True)

table = add_table_zh(doc, 4, 4)
headers = ['服务商', '免费额度', '特点', '推荐度']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

db_providers = [
    ['Supabase', '500MB存储, 2GB传输', '功能最全，含Auth/Storage/Realtime', '★★★★★'],
    ['Vercel Postgres', '256MB存储', '与Vercel无缝集成', '★★★★'],
    ['Neon', '3GB存储', 'Serverless，无连接限制', '★★★★']
]
for i, row_data in enumerate(db_providers, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 9)

add_heading_zh(doc, '6.4 部署后操作', 2)
after_deploy = [
    '创建数据库实例',
    '执行Prisma迁移（npx prisma migrate deploy）',
    '初始化商品数据',
    '配置环境变量到Vercel',
    '测试API连接'
]
for i, step in enumerate(after_deploy, 1):
    add_paragraph_zh(doc, f'{i}. {step}')

doc.add_page_break()

# ========== 第7章：常见问题与故障处理 ==========
add_heading_zh(doc, '7. 常见问题与故障处理', 1)

add_heading_zh(doc, '7.1 支付问题', 2)

payment_issues = [
    ['客户付错金额', '联系客户补款或退款重付'],
    ['客户付错网络', '尝试找回，如无法找回协商处理'],
    ['长时间未到账', '检查TXID是否正确，联系客户确认'],
    ['客户说已付但查不到', '索要TXID和付款截图，链上核实']
]
table = add_table_zh(doc, len(payment_issues) + 1, 2)
headers = ['问题', '解决方案']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

for i, row_data in enumerate(payment_issues, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 10)

add_heading_zh(doc, '7.2 系统问题', 2)

system_issues = [
    ['无法登录后台', '检查用户名密码，清除浏览器缓存'],
    ['页面显示异常', '强制刷新（Ctrl+F5），检查网络'],
    ['API接口报错', '检查数据库连接，查看Vercel日志'],
    ['网站无法访问', '检查Surge.sh状态，确认域名解析']
]
table = add_table_zh(doc, len(system_issues) + 1, 2)
headers = ['问题', '解决方案']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

for i, row_data in enumerate(system_issues, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 10)

add_heading_zh(doc, '7.3 紧急情况处理', 2)

table = add_table_zh(doc, 5, 2)
headers = ['情况', '处理方案']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

emergency = [
    ['L1钱包异常', '立即启用备用地址，更新网站配置'],
    ['L2资金不足', '从L1紧急调入，记录后补'],
    ['大量订单涌入', '暂停广告投放，优先处理已有订单'],
    ['客户集中投诉', '暂停销售，排查问题，逐个解决'],
    ['网站被攻击', '联系Vercel支持，启用DDoS防护']
]
table = add_table_zh(doc, len(emergency) + 1, 2)
headers = ['情况', '处理方案']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    run = cell.paragraphs[0].add_run(header)
    set_chinese_font(run, 'SimHei', 11, True)

for i, row_data in enumerate(emergency, 1):
    for j, cell_data in enumerate(row_data):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(cell_data)
        set_chinese_font(run, 'SimSun', 10)

doc.add_page_break()

# ========== 第8章：附录 ==========
add_heading_zh(doc, '8. 附录', 1)

add_heading_zh(doc, '8.1 常用网址清单', 2)
urls_list = [
    '主站：https://horizon-watch-store-1773050228.surge.sh',
    '后台：https://horizon-watch-store-1773050228.surge.sh/admin',
    'GitHub：https://github.com/999huamao-ctrl/vintage-watch-co',
    'Vercel Dashboard：https://vercel.com/dashboard',
    'Supabase：https://supabase.com',
    'TronScan（查账）：https://tronscan.org'
]
for url in urls_list:
    add_paragraph_zh(doc, '• ' + url)

add_heading_zh(doc, '8.2 联系方式', 2)
contacts = [
    '投资人：STAR (999huamao@gmail.com)',
    'CEO/运营：S1 (自动系统)',
    '日报发送：每日22:00自动发送'
]
for c in contacts:
    add_paragraph_zh(doc, '• ' + c)

add_heading_zh(doc, '8.3 更新日志', 2)
changelog = [
    'v2.2 (2026-03-16) - 统一Vercel域名、更新数据库状态说明',
    'v2.1 (2026-03-16) - 更新生产环境URL、修正版本号',
    'v2.0 (2026-03-15) - 新增3级钱包架构、用户管理功能、数据库Schema',
    'v2.1 (2026-03-11) - 新增AI客服、多角色权限、USDT收款流程',
    'v2.0 (2026-03-08) - UI改版、竞品调研、广告方案',
    'v1.0 (2026-03-05) - 网站基础框架、商品展示、支付集成'
]
for log in changelog:
    add_paragraph_zh(doc, '• ' + log)

# 保存文档
output_path = '/root/.openclaw/workspace/projects/watch-store/docs/Horizon_Watches_运营手册_v2.2.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
