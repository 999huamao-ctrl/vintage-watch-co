from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

def set_chinese_font(run, font_name='SimSun', size=11, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    if level == 1:
        set_chinese_font(run, 'SimHei', 16, True)
        run.font.color.rgb = RGBColor(0, 0, 128)
    elif level == 2:
        set_chinese_font(run, 'SimHei', 14, True)
    else:
        set_chinese_font(run, 'SimHei', 12, True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, 'SimSun', size, bold)
    return p

def add_table_zh(doc, rows, cols, style='Table Grid'):
    """添加表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    return table

# 创建文档
doc = Document()

# 设置文档默认字体
doc.styles['Normal'].font.name = 'SimSun'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# ========== 封面 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Horizon Watches\n运营手册 v3.1')
set_chinese_font(run, 'SimHei', 28, True)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
now = datetime.now()
date_time_str = now.strftime('版本 3.1 | %Y年%m月%d日 %H:%M 生成')
run = subtitle.add_run(date_time_str)
set_chinese_font(run, 'SimSun', 14)

doc.add_page_break()

# ========== 第1章：系统访问信息 ==========
add_heading_zh(doc, '1. 系统访问信息', 1)

add_heading_zh(doc, '1.1 网站地址', 2)
table = add_table_zh(doc, 3, 2)
table.rows[0].cells[0].text = '项目'
table.rows[0].cells[1].text = '访问地址'
table.rows[1].cells[0].text = '主站 (客户访问)'
table.rows[1].cells[1].text = 'https://www.horizonoo.cc'
table.rows[2].cells[0].text = '管理后台'
table.rows[2].cells[1].text = 'https://www.horizonoo.cc/admin'

add_heading_zh(doc, '1.2 后台登录', 2)
add_paragraph_zh(doc, '1. 访问 https://www.horizonoo.cc/admin')
add_paragraph_zh(doc, '2. 输入用户名和密码')
add_paragraph_zh(doc, '3. 根据角色权限进入对应仪表盘')
add_paragraph_zh(doc, '注意：首次登录请使用分配的账号密码，登录后建议修改密码', bold=True)

doc.add_page_break()

# ========== 第2章：角色权限说明 ==========
add_heading_zh(doc, '2. 角色权限说明', 1)

table = add_table_zh(doc, 5, 2)
table.rows[0].cells[0].text = '角色'
table.rows[0].cells[1].text = '功能权限'
table.rows[1].cells[0].text = 'SUPERADMIN'
table.rows[1].cells[1].text = '全权限 (含钱包设置、用户管理)'
table.rows[2].cells[0].text = 'ADMIN'
table.rows[2].cells[1].text = '商品管理、订单管理、库存管理'
table.rows[3].cells[0].text = 'SUPPLY'
table.rows[3].cells[1].text = '仅库存管理 (修改库存数量)'
table.rows[4].cells[0].text = 'LOGISTICS'
table.rows[4].cells[1].text = '仅订单发货 (更新物流信息)'

doc.add_page_break()

# ========== 第3章：商品数据导入指南 ==========
add_heading_zh(doc, '3. 商品数据导入指南', 1)

add_heading_zh(doc, '3.1 支持的导入方式', 2)

add_paragraph_zh(doc, '方式1: 后台手动添加', bold=True)
add_paragraph_zh(doc, '• 适合少量商品 (1-5个)')
add_paragraph_zh(doc, '• 进入后台 → Products → Add Product')
add_paragraph_zh(doc, '• 逐个填写商品信息')

add_paragraph_zh(doc, '方式2: CSV批量导入', bold=True)
add_paragraph_zh(doc, '• 适合批量商品更新')
add_paragraph_zh(doc, '• 发送表格给技术团队导入')
add_paragraph_zh(doc, '• 或直接在后台使用导入功能')

add_heading_zh(doc, '3.2 CSV表格格式要求', 2)
add_paragraph_zh(doc, '必需字段 (缺少则无法导入):', bold=True)

table = add_table_zh(doc, 6, 3)
table.rows[0].cells[0].text = '字段名'
table.rows[0].cells[1].text = '说明'
table.rows[0].cells[2].text = '示例'
table.rows[1].cells[0].text = 'name'
table.rows[1].cells[1].text = '商品名称 (英文)'
table.rows[1].cells[2].text = 'Rolex Submariner Style'
table.rows[2].cells[0].text = 'price'
table.rows[2].cells[1].text = '售价 (€)'
table.rows[2].cells[2].text = '189.00'
table.rows[3].cells[0].text = 'category'
table.rows[3].cells[1].text = '分类'
table.rows[3].cells[2].text = 'Men / Women'
table.rows[4].cells[0].text = 'stock'
table.rows[4].cells[1].text = '库存数量'
table.rows[4].cells[2].text = '15'
table.rows[5].cells[0].text = 'image'
table.rows[5].cells[1].text = '主图URL'
table.rows[5].cells[2].text = '/products/image1.webp'

add_paragraph_zh(doc, '可选字段:', bold=True)

table = add_table_zh(doc, 14, 3)
table.rows[0].cells[0].text = '字段名'
table.rows[0].cells[1].text = '说明'
table.rows[0].cells[2].text = '示例'
table.rows[1].cells[0].text = 'nameZh'
table.rows[1].cells[1].text = '中文名称'
table.rows[1].cells[2].text = '劳力士潜航者风格'
table.rows[2].cells[0].text = 'nameDe'
table.rows[2].cells[1].text = '德语名称'
table.rows[2].cells[2].text = 'Rolex Submariner Stil'
table.rows[3].cells[0].text = 'nameFr'
table.rows[3].cells[1].text = '法语名称'
table.rows[3].cells[2].text = 'Style Rolex Submariner'
table.rows[4].cells[0].text = 'nameEs'
table.rows[4].cells[1].text = '西班牙语名称'
table.rows[4].cells[2].text = 'Estilo Rolex Submariner'
table.rows[5].cells[0].text = 'nameIt'
table.rows[5].cells[1].text = '意大利语名称'
table.rows[5].cells[2].text = 'Stile Rolex Submariner'
table.rows[6].cells[0].text = 'nameJa'
table.rows[6].cells[1].text = '日语名称'
table.rows[6].cells[2].text = 'ロレックス サブマリーナ スタイル'
table.rows[7].cells[0].text = 'originalPrice'
table.rows[7].cells[1].text = '原价/划线价 (€)'
table.rows[7].cells[2].text = '299.00'
table.rows[8].cells[0].text = 'description'
table.rows[8].cells[1].text = '商品描述'
table.rows[8].cells[2].text = 'Classic diving watch...'
table.rows[9].cells[0].text = 'detailImage1-4'
table.rows[9].cells[1].text = '细节图1-4 URL'
table.rows[9].cells[2].text = '/products/detail1.webp'
table.rows[10].cells[0].text = 'caseSize'
table.rows[10].cells[1].text = '表壳尺寸'
table.rows[10].cells[2].text = '40mm'
table.rows[11].cells[0].text = 'movement'
table.rows[11].cells[1].text = '机芯类型'
table.rows[11].cells[2].text = 'Automatic'
table.rows[12].cells[0].text = 'strap'
table.rows[12].cells[1].text = '表带材质'
table.rows[12].cells[2].text = 'Stainless Steel'
table.rows[13].cells[0].text = 'waterResistance'
table.rows[13].cells[1].text = '防水等级'
table.rows[13].cells[2].text = '300m'

add_heading_zh(doc, '3.3 示例表格格式', 2)
add_paragraph_zh(doc, 'CSV文件内容示例：')
add_paragraph_zh(doc, 'name,nameZh,price,originalPrice,category,stock,image,caseSize,movement,strap,waterResistance')
add_paragraph_zh(doc, 'Rolex Submariner Style,劳力士潜航者风格,189.00,299.00,Men,15,/products/submariner.webp,40mm,Automatic,Stainless Steel,300m')
add_paragraph_zh(doc, 'Rolex Datejust Style,劳力士日志型风格,169.00,259.00,Men,12,/products/datejust.webp,36mm,Automatic,Jubilee Bracelet,100m')
add_paragraph_zh(doc, 'Rolex Lady-Datejust Style,劳力士女装日志型风格,169.00,269.00,Women,18,/products/lady-datejust.webp,28mm,Automatic,Jubilee Bracelet,100m')

add_heading_zh(doc, '3.4 常见问题', 2)

add_paragraph_zh(doc, 'Q: 图片URL怎么填?', bold=True)
add_paragraph_zh(doc, 'A: 图片需要先上传到图床或CDN，填入完整的图片访问地址。格式如:')
add_paragraph_zh(doc, '• 相对路径: /products/image1.webp')
add_paragraph_zh(doc, '• 完整URL: https://cdn.horizonoo.cc/products/image1.webp')

add_paragraph_zh(doc, 'Q: 价格格式有什么要求?', bold=True)
add_paragraph_zh(doc, 'A: 使用数字格式，不要加货币符号。例如:')
add_paragraph_zh(doc, '• 正确: 189.00 或 189')
add_paragraph_zh(doc, '• 错误: €189.00 或 189€')

add_paragraph_zh(doc, 'Q: 如何下架商品?', bold=True)
add_paragraph_zh(doc, 'A: 将 isActive 字段设为 false，或直接在后台关闭商品')

add_paragraph_zh(doc, 'Q: 更新现有商品数据?', bold=True)
add_paragraph_zh(doc, 'A: 在表格中保留商品的 id 字段，系统会根据id更新对应商品。如没有id字段，则会创建新商品')

doc.add_page_break()

# ========== 第4章：快速操作指南 ==========
add_heading_zh(doc, '4. 快速操作指南', 1)

add_heading_zh(doc, '4.1 更新商品库存 (SUPPLY角色)', 2)
add_paragraph_zh(doc, '1. 登录后台 → Products')
add_paragraph_zh(doc, '2. 找到需要更新的商品')
add_paragraph_zh(doc, '3. 点击 Edit')
add_paragraph_zh(doc, '4. 修改 Stock 数量')
add_paragraph_zh(doc, '5. 点击 Save')

add_heading_zh(doc, '4.2 处理订单发货 (LOGISTICS角色)', 2)
add_paragraph_zh(doc, '1. 登录后台 → Orders')
add_paragraph_zh(doc, '2. 找到 PAID 状态的订单')
add_paragraph_zh(doc, '3. 点击订单查看详情')
add_paragraph_zh(doc, '4. 更新状态为 SHIPPED')
add_paragraph_zh(doc, '5. 填写 Tracking Number (物流单号)')
add_paragraph_zh(doc, '6. 填写 Carrier (承运商)')
add_paragraph_zh(doc, '7. 点击 Update')

add_heading_zh(doc, '4.3 添加新商品 (ADMIN角色)', 2)
add_paragraph_zh(doc, '1. 登录后台 → Products')
add_paragraph_zh(doc, '2. 点击 Add Product')
add_paragraph_zh(doc, '3. 填写所有商品信息')
add_paragraph_zh(doc, '4. 点击 Save')

doc.add_page_break()

# ========== 第5章：注意事项 ==========
add_heading_zh(doc, '5. 注意事项', 1)

add_paragraph_zh(doc, '1. 商品名称: 建议使用英文，多语言版本填入对应字段')
add_paragraph_zh(doc, '2. 分类: 目前仅支持 Men 或 Women')
add_paragraph_zh(doc, '3. 库存: 设为0时商品不会下架，但会显示"缺货"')
add_paragraph_zh(doc, '4. 图片: 建议尺寸 800x800px 以上，格式 webp 或 jpg')
add_paragraph_zh(doc, '5. 价格: 建议同时填写 price 和 originalPrice 以显示折扣效果')

doc.add_page_break()

# ========== 第6章：技术支持 ==========
add_heading_zh(doc, '6. 技术支持', 1)

add_paragraph_zh(doc, '如有问题请联系技术团队:')
add_paragraph_zh(doc, '• 系统故障/Bug 报告')
add_paragraph_zh(doc, '• 数据导入协助')
add_paragraph_zh(doc, '• 功能使用咨询')

add_paragraph_zh(doc, '')
add_paragraph_zh(doc, '---')
add_paragraph_zh(doc, '文档版本: v3.1')
add_paragraph_zh(doc, '更新日期: 2026-03-23')
add_paragraph_zh(doc, '更新内容: 新增商品CSV导入格式说明，更新后台访问地址')

# 保存文档
output_path = '/root/.openclaw/workspace/projects/watch-store/docs/Horizon_Watches_运营手册_v3.1.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
