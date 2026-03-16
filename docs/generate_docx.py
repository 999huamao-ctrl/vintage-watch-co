#!/usr/bin/env python3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# 标题
title = doc.add_heading('Horizon Watches 运营手册', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 元信息
doc.add_paragraph('版本: v2.0 | 更新日期: 2026-03-15 | 编制: S1')
doc.add_paragraph()

# 一、项目概况
doc.add_heading('一、项目概况', 1)
doc.add_heading('1.1 项目定位', 2)
table = doc.add_table(rows=7, cols=2)
table.style = 'Light Grid Accent 1'
cells = [
    ('维度', '内容'),
    ('品牌名', 'Horizon Watches'),
    ('定位', '欧洲独立站D2C手表电商'),
    ('目标市场', '德国、法国、意大利、西班牙等17国'),
    ('产品', '14款劳力士风格腕表'),
    ('客单价', '€69-89'),
    ('毛利率', '平均50.6%'),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('1.2 核心资源', 2)
doc.add_paragraph('供应链: ✓ 已确认（STAR提供）', style='List Bullet')
doc.add_paragraph('物流渠道: ✓ 欧洲敏感专线，9-15工作日', style='List Bullet')
doc.add_paragraph('技术栈: Next.js + Vercel + Prisma + Supabase', style='List Bullet')
doc.add_paragraph('支付方式: PayPal + USDT(TRC20)', style='List Bullet')

# 二、网站运营
doc.add_heading('二、网站运营', 1)
doc.add_heading('2.1 网站地址', 2)
doc.add_paragraph('生产环境: https://frontend-bxupmyqv8-999huamao-ctrls-projects.vercel.app')
doc.add_paragraph('管理后台: /admin')

doc.add_heading('2.2 后台账号体系', 2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['账号', '密码', '角色', '权限']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('superadmin', 'super123', 'SuperAdmin', '全权限（含钱包设置）'),
    ('admin', 'admin123', 'Admin', '产品、订单、财务、分析'),
    ('supply', 'supply123', 'Supply', '仅商品管理'),
    ('logistics', 'logistics123', 'Logistics', '仅订单发货'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('2.3 3级钱包架构', 2)
doc.add_paragraph('L1收款钱包 → L2运营钱包 → L3利润钱包')
doc.add_paragraph('(客户付款)  (40%广告/物流)  (60%利润归STAR)')
doc.add_paragraph()
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['层级', '地址', '用途']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('L1', 'TGPBhfjSuwjrfGUtdqt6EZUbzhbRCGfC5c', '仅接收客户USDT付款'),
    ('L2', 'TCWgr2qGcheRsD7kceoFpJfMg59fFrJGCq', '广告、物流、运营成本'),
    ('L3', 'TUyTqV47pd7o3Bg6Uhw5XJ9rwkdgi6tsKb', '最终利润归属'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('资金流转规则:', style='List Bullet')
doc.add_paragraph('每日归集: L1 → L2(40%) + L3(60%)', style='List Bullet')
doc.add_paragraph('L2余额低于€50时预警，需补充', style='List Bullet')

# 三、产品管理
doc.add_heading('三、产品管理', 1)
doc.add_heading('3.1 产品目录（14款）', 2)
table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['SKU', '名称', '售价€', '进货¥', '毛利']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('H001', 'Submariner Vintage', '89', '420', '52.8%'),
    ('H002', 'GMT Master II', '95', '480', '49.5%'),
    ('H003', 'Daytona Classic', '99', '520', '47.5%'),
    ('...', '共14款', '-', '-', '平均50.6%'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('3.2 库存管理', 2)
doc.add_paragraph('后台操作: Supply角色可编辑库存', style='List Bullet')
doc.add_paragraph('低库存预警: <5件标红提醒', style='List Bullet')
doc.add_paragraph('补货流程: 库存归零前48小时通知STAR', style='List Bullet')

# 四、订单处理流程
doc.add_heading('四、订单处理流程', 1)
doc.add_heading('4.1 订单状态流转', 2)
doc.add_paragraph('PENDING → PAID → PROCESSING → SHIPPED → DELIVERED')
doc.add_paragraph('  └────── CANCELLED (超时未支付)')

doc.add_heading('4.2 各角色职责', 2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
headers = ['角色', '职责']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('系统', '自动确认PayPal/USDT支付，状态→PAID'),
    ('Logistics', '收到PAID订单后24小时内发货'),
    ('Logistics', '回填物流单号，状态→SHIPPED'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('4.3 物流信息', 2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
cells = [
    ('项目', '内容'),
    ('渠道', '欧洲敏感专线-欧美-通'),
    ('时效', '9-15工作日'),
    ('覆盖', '德/法/意/西/葡/荷/比/奥/瑞/丹/挪/芬/波/捷/英'),
    ('成本', '€5-8/只(0.3-0.5kg)'),
    ('追踪', '提供追踪单号，可在17track查询'),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# 五、广告投放运营
doc.add_heading('五、广告投放运营', 1)
doc.add_heading('5.1 启动配置', 2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
cells = [
    ('项目', '配置'),
    ('启动预算', '€100'),
    ('投放平台', 'Facebook/Instagram'),
    ('目标国家', '德国、法国、意大利（首期）'),
    ('出价策略', '最低成本(Cost Cap €15/转化)'),
    ('归因窗口', '7天点击+1天浏览'),
]
for i, (k, v) in enumerate(cells):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('5.2 广告素材规范', 2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['版位', '尺寸', '用途']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('Feed', '1080×1080', '主力转化广告'),
    ('Stories', '1080×1920', '品牌曝光'),
    ('Link', '1200×628', '链接点击优化'),
    ('Carousel', '1080×1080', '多产品展示'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('5.3 再投资策略', 2)
doc.add_paragraph('月收入 × 40% = 下月广告预算')
doc.add_paragraph('例: 月入€1000 → 下月预算€400')

doc.add_heading('5.4 投放时间表', 2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['阶段', '时间', '动作', '预算']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('Day 1', '3月17日', 'CBO测试3组素材', '€30'),
    ('Day 3', '3月19日', '关停CTR<1%，保留胜者', '-'),
    ('Day 7', '3月23日', '规模化胜出素材', '€70'),
    ('Day 14', '3月30日', '分析ROI，调整受众', '动态'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 六、AI客服配置
doc.add_heading('六、AI客服配置', 1)
doc.add_heading('6.1 知识库覆盖', 2)
doc.add_paragraph('14款产品详情（价格、规格、卖点）', style='List Bullet')
doc.add_paragraph('物流时效与追踪', style='List Bullet')
doc.add_paragraph('PayPal/USDT支付流程', style='List Bullet')
doc.add_paragraph('7天无理由退换政策', style='List Bullet')

doc.add_heading('6.2 快捷回复按钮', 2)
doc.add_paragraph('「产品价格」「配送时效」「支付方式」「转人工」')

doc.add_heading('6.3 转人工触发条件', 2)
doc.add_paragraph('客户连续3次未被匹配到答案', style='List Bullet')
doc.add_paragraph('客户输入"人工"/"客服"/"human"', style='List Bullet')
doc.add_paragraph('退款/投诉类关键词', style='List Bullet')

# 七、数据监控
doc.add_heading('七、数据监控', 1)
doc.add_heading('7.1 关键指标看板', 2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['指标', '目标', '监控频率']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('日访客(UV)', '>200', '每日'),
    ('转化率', '>2%', '每日'),
    ('客单价', '€75+', '每周'),
    ('ROAS', '>2.0', '每周'),
    ('退货率', '<5%', '每月'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('7.2 监控工具', 2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
headers = ['工具', '用途']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('Vercel Analytics', '网站性能、Core Web Vitals'),
    ('Facebook Pixel', '广告转化追踪'),
    ('Google Analytics 4', '流量来源分析'),
    ('后台Analytics页', '订单/收入/转化数据'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 八、风险管理
doc.add_heading('八、风险管理', 1)
doc.add_heading('8.1 支付风险', 2)
doc.add_paragraph('PayPal新账户限额: €2500/月', style='List Bullet')
doc.add_paragraph('达到80%限额时预警，准备备用账户', style='List Bullet')
doc.add_paragraph('USDT收款作为备用通道', style='List Bullet')

doc.add_heading('8.2 广告账户风险', 2)
doc.add_paragraph('新账户首周日预算上限€50', style='List Bullet')
doc.add_paragraph('避免短时间内大幅调整出价', style='List Bullet')
doc.add_paragraph('准备多个BM和账户备用', style='List Bullet')

doc.add_heading('8.3 物流风险', 2)
doc.add_paragraph('敏感货渠道可能查验', style='List Bullet')
doc.add_paragraph('客户投诉处理: 48小时内响应', style='List Bullet')
doc.add_paragraph('丢件赔付: 全款退或重发', style='List Bullet')

# 九、待决策事项
doc.add_heading('九、待决策事项', 1)
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['序号', '事项', '状态', '说明']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('1', 'Facebook广告充值方式', '待决策', '虚拟卡(Depay/OneKey) vs 代理商代投'),
    ('2', '数据库服务商选择', '待决策', 'Vercel Postgres / Supabase / Neon'),
    ('3', '首批广告启动时间', '待确认', '建议3月17日(周一)'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 十、联系信息
doc.add_heading('十、联系信息', 1)
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['角色', '联系方式', '用途']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('客服邮箱', 'support@horizonwatches.com', '客户支持'),
    ('投资人', 'STAR', '资金/战略决策'),
    ('CEO', 'S1', '日常运营/技术'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph()

# 版本记录
doc.add_heading('文档版本记录', 1)
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['版本', '日期', '更新内容']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ('v1.0', '2026-03-07', '初始版本'),
    ('v2.0', '2026-03-15', '增加数据库、后台权限、系统状态监控'),
]
for i, row in enumerate(data, 1):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 保存
doc.save('/root/.openclaw/workspace/projects/watch-store/docs/OPERATIONS_MANUAL_v2.docx')
print('Word文档已生成: OPERATIONS_MANUAL_v2.docx')
